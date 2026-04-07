# src/document_processor.py
import os
import fitz  # PyMuPDF for PDF
from docx import Document # python-docx for DOCX
import openpyxl # openpyxl for XLSX
from PIL import Image # Pillow for image manipulation
# import pytesseract # pytesseract for OCR
from typing import List, Tuple, Dict, Union
from io import BytesIO
from pathlib import Path

# Import paths from utils
from src.utils import RAW_DOCUMENTS_DIR

# --- OCR Configuration (for images) ---
# TESSERACT_CMD_PATH: Set this if tesseract is not in your system's PATH.
# Example for Windows: r'C:\Program Files\Tesseract-OCR\tesseract.exe'
# Example for Linux/macOS (if not in PATH): '/usr/local/bin/tesseract'
# If Tesseract is in your system's PATH, you can leave this as None.

print("Please install Tesseract OCR and/or set TESSERACT_CMD_PATH in document_processor.py.")


def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extracts text from a PDF file using PyMuPDF."""
    text = ""
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            text += page.get_text()
        doc.close()
        print(f"Extracted text from PDF: {pdf_path.name}")
    except fitz.FileDataError as e:
        print(f"Error extracting text from PDF {pdf_path.name}: {e}")
    except Exception as e:
        print(f"An unexpected error occurred while processing PDF {pdf_path.name}: {e}")
    return text

def extract_text_from_docx(docx_path: Path) -> str:
    """Extracts text from a DOCX file using python-docx."""
    text = ""
    try:
        doc = Document(docx_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        print(f"Extracted text from DOCX: {docx_path.name}")
    except Exception as e:
        print(f"Error extracting text from DOCX {docx_path.name}: {e}")
    return text

def extract_text_from_xlsx(xlsx_path: Path) -> str:
    """Extracts text from an XLSX file using openpyxl."""
    text = ""
    try:
        workbook = openpyxl.load_workbook(xlsx_path)
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text += f"--- Sheet: {sheet_name} ---\n"
            for row in sheet.iter_rows():
                row_values = [str(cell.value) if cell.value is not None else "" for cell in row]
                text += "\t".join(row_values) + "\n"
            text += "\n" # Add a newline between sheets
        print(f"Extracted text from XLSX: {xlsx_path.name}")
    except Exception as e:
        print(f"Error extracting text from XLSX {xlsx_path.name}: {e}")
    return text

def extract_text_from_image(image_path: Path) -> str:
    """
    Extracts text from an image file.
    NOTE: OCR is disabled for deployment (Streamlit Cloud doesn't support Tesseract).
    """
    try:
        img = Image.open(image_path)
        print(f"OCR skipped for image: {image_path.name}")
        return ""
    except Exception as e:
        print(f"Error opening image {image_path.name}: {e}")
        return ""

def process_documents(directory_path: Path = RAW_DOCUMENTS_DIR) -> List[Tuple[str, Dict]]:
    """
    Processes all supported documents in a given directory, extracts text,
    and returns a list of (text_content, metadata) tuples.
    """
    processed_data = []
    supported_extensions = {
        ".pdf": extract_text_from_pdf,
        ".docx": extract_text_from_docx,
        ".xlsx": extract_text_from_xlsx,
        ".png": extract_text_from_image,
        ".jpg": extract_text_from_image,
        ".jpeg": extract_text_from_image,
    }

    print(f"Starting document processing in: {directory_path}")
    if not directory_path.exists():
        print(f"Error: Directory not found at {directory_path}")
        return []

    for file_name in os.listdir(directory_path):
        file_path = directory_path / file_name
        if file_path.is_file():
            file_extension = file_path.suffix.lower()
            
            if file_extension in supported_extensions:
                print(f"Processing {file_name}...")
                extractor_func = supported_extensions[file_extension]
                text_content = extractor_func(file_path)
                
                if text_content:
                    metadata = {
                        "file_name": file_name,
                        "file_extension": file_extension,
                        "source_path": str(file_path)  # Convert Path to string
                    }
                    processed_data.append((text_content, metadata))
                else:
                    print(f"No text extracted or error for {file_name}. Skipping.")
            else:
                print(f"Skipping unsupported file type: {file_name}")
    
    print(f"Finished document processing. Total documents processed: {len(processed_data)}")
    return processed_data

# --- For testing the document_processor directly ---
if __name__ == '__main__':
    print("--- Testing Document Processor ---")

    # Create dummy raw documents for testing
    print(f"Ensuring raw documents directory exists: {RAW_DOCUMENTS_DIR}")
    RAW_DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)

    dummy_pdf_path = RAW_DOCUMENTS_DIR / "dummy_admission_guide.pdf"
    dummy_docx_path = RAW_DOCUMENTS_DIR / "dummy_faq.docx"
    dummy_xlsx_path = RAW_DOCUMENTS_DIR / "dummy_fees.xlsx"
    dummy_image_path = RAW_DOCUMENTS_DIR / "dummy_text_image.png"

    # Create dummy PDF (requires PyMuPDF to write, but we can fake it for testing)
    try:
        # Create a tiny PDF file directly for robust testing
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "This is a sample admission guide. Apply by Dec 15.", fontname="helv", fontsize=12)
        doc.save(dummy_pdf_path)
        doc.close()
        print(f"Created dummy PDF: {dummy_pdf_path.name}")
    except Exception as e:
        print(f"Could not create dummy PDF (requires PyMuPDF): {e}")
        # If PyMuPDF can't write, ensure the file is clean
        if dummy_pdf_path.exists(): os.remove(dummy_pdf_path)


    # Create dummy DOCX
    try:
        document = Document()
        document.add_heading('Admission FAQs', level=1)
        document.add_paragraph('Q: What is the minimum GPA?')
        document.add_paragraph('A: A minimum GPA of 3.0 is required.')
        document.add_paragraph('Q: Are scholarships available?')
        document.add_paragraph('A: Yes, scholarships are available for deserving students.')
        document.save(dummy_docx_path)
        print(f"Created dummy DOCX: {dummy_docx_path.name}")
    except Exception as e:
        print(f"Could not create dummy DOCX (requires python-docx): {e}")
        if dummy_docx_path.exists(): os.remove(dummy_docx_path)

    # Create dummy XLSX
    try:
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "Fees 2024-2025"
        sheet['A1'] = "Program"
        sheet['B1'] = "Annual Fee"
        sheet['A2'] = "B.Tech"
        sheet['B2'] = "1,50,000 INR"
        sheet['A3'] = "M.Tech"
        sheet['B3'] = "1,00,000 INR"
        workbook.save(dummy_xlsx_path)
        print(f"Created dummy XLSX: {dummy_xlsx_path.name}")
    except Exception as e:
        print(f"Could not create dummy XLSX (requires openpyxl): {e}")
        if dummy_xlsx_path.exists(): os.remove(dummy_xlsx_path)

    # Create dummy Image with text (requires Pillow and assumes Tesseract is installed for OCR)
    try:
        from PIL import Image, ImageDraw, ImageFont
        # Create a blank image with some text
        img_width, img_height = 400, 100
        img = Image.new('RGB', (img_width, img_height), color = (255, 255, 255))
        d = ImageDraw.Draw(img)
        # Try to use a common font or default
        try:
            # Adjust font path for your OS if necessary.
            # Example for Windows: "arial.ttf"
            # Example for Linux: "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
            # Example for macOS: "/System/Library/Fonts/SFCompactText-Regular.otf"
            font = ImageFont.truetype("arial.ttf", 20) # Try a common font
        except IOError:
            font = ImageFont.load_default() # Fallback to default font
            print("Warning: Could not load 'arial.ttf', using default font for image creation.")

        text_to_draw = "Welcome to University!"
        text_bbox = d.textbbox((0,0), text_to_draw, font=font)
        text_width = text_bbox[2] - text_bbox[0]
        text_height = text_bbox[3] - text_bbox[1]
        
        # Center the text
        x = (img_width - text_width) / 2
        y = (img_height - text_height) / 2
        d.text((x, y), text_to_draw, fill=(0,0,0), font=font)
        img.save(dummy_image_path)
        print(f"Created dummy Image: {dummy_image_path.name}")
    except Exception as e:
        print(f"Could not create dummy Image (requires Pillow): {e}")
        if dummy_image_path.exists(): os.remove(dummy_image_path)


    # Process the dummy documents
    print("\n--- Running document processing ---")
    processed_docs = process_documents(RAW_DOCUMENTS_DIR)

    if processed_docs:
        for i, (text, meta) in enumerate(processed_docs):
            print(f"\nDocument {i+1}:")
            print(f"  File: {meta['file_name']}")
            print(f"  Extension: {meta['file_extension']}")
            print(f"  Content (first 200 chars): {text[:200]}...")
    else:
        print("No documents were processed. Check for errors or if dummy files were created.")

    # Clean up dummy files
    print("\n--- Cleaning up dummy files ---")
    for f_path in [dummy_pdf_path, dummy_docx_path, dummy_xlsx_path, dummy_image_path]:
        if f_path.exists():
            try:
                os.remove(f_path)
                print(f"Removed: {f_path.name}")
            except Exception as e:
                print(f"Error removing {f_path.name}: {e}")

    print("\n--- Document Processor tests complete ---")